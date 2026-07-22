"""
Triagem de candidato por IA, sob demanda.

E o item que a origem trata como o coracao do produto: "como se voce tivesse um
recrutador trabalhando por voce sem voce nem estar dentro da plataforma".

TRES DECISOES QUE MOLDAM ESTE MODULO:

1. SOB DEMANDA. Roda quando o RH aperta "Analisar", e nao na chegada da
   candidatura. Nao e detalhe de implementacao: rodar na chegada faria o
   candidato esperar a LLM no celular, e faria o tenant pagar analise dos 76
   que chegaram, inclusive dos que ninguem vai olhar.

2. NUNCA MOVE O CANDIDATO. Devolve sugestao; a decisao continua humana. A
   origem repete isso em dois lugares ("sempre precisa de revisao humana"), e
   automatizar recusa seria decidir contratacao por LLM sem ninguem pra
   responder por ela.

3. SO USA O QUE A VAGA DECLAROU. O criterio sai de `requisitos_de_triagem()`,
   que o RH cadastrou. Sem requisito cadastrado nao ha triagem: a IA inventaria
   criterio proprio, e ai a "consistencia" que ela promete vira arbitrariedade
   com cara de objetividade.

O QUE NAO VAI PRO PROMPT, de proposito: nome, WhatsApp, email e endereco. Eles
nao ajudam a avaliar aptidao e so aumentam a exposicao de dado pessoal a um
terceiro. Cidade e bairro entram porque deslocamento e criterio real de vaga
operacional. Ver `_perfil_do_candidato`.
"""
import json
import logging

from apps.automacao.services.ia import chamar_llm, integracao_ia_do_tenant
from apps.people.excecoes import PeopleError
from apps.people.models import AnaliseCandidato

logger = logging.getLogger(__name__)

LIMITE_CURRICULO = 12000   # caracteres. CV de 3 paginas cabe folgado.


class TriagemIndisponivel(PeopleError):
    """Falta integracao de IA, ou falta o que analisar."""


def _texto_de_pdf(arquivo):
    from pypdf import PdfReader

    leitor = PdfReader(arquivo)
    return '\n'.join((pagina.extract_text() or '') for pagina in leitor.pages)


def _texto_de_docx(arquivo):
    """
    Paragrafos E tabelas.

    As tabelas nao sao detalhe: modelo de curriculo do Word usa tabela pra
    diagramar, e e comum o historico profissional inteiro morar dentro de uma.
    Ler so `doc.paragraphs` devolveria o nome da pessoa e mais nada, e a analise
    sairia dizendo "informou pouco" sobre um curriculo completo.
    """
    import docx

    doc = docx.Document(arquivo)
    pedacos = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            pedacos.extend(celula.text for celula in linha.cells)
    return '\n'.join(pedacos)


# Extensao -> extrator. Mapa, e nao if/elif, porque e o mesmo criterio que
# `EXTENSOES_CURRICULO` usa pra decidir o que aceitar no upload: os dois precisam
# andar juntos, e um teste amarra isso.
LEITORES = {'.pdf': _texto_de_pdf, '.docx': _texto_de_docx}


def _texto_do_curriculo(candidato):
    """
    Texto do curriculo, ou vazio.

    Devolve vazio em vez de levantar: curriculo ilegivel nao impede a analise,
    so a deixa mais pobre, e isso e registrado em `usou_curriculo` pra o RH saber
    que a sugestao saiu so do formulario.

    Imagem nao e lida: precisaria de OCR, que e dependencia bem mais pesada. O
    candidato que manda foto continua sendo analisado pelo que preencheu.

    `.doc` (Word binario, pre 2007) tambem nao: nenhuma biblioteca Python leve
    le. Por isso ele saiu tambem da lista de formatos ACEITOS, em vez de ficar
    aceito e ilegivel. Prometer o que nao se le e o que produz analise cega com
    cara de completa.
    """
    if not candidato.curriculo:
        return ''

    nome = candidato.curriculo.name.lower()
    leitor = next((f for ext, f in LEITORES.items() if nome.endswith(ext)), None)
    if leitor is None:
        return ''

    try:
        candidato.curriculo.open('rb')
        try:
            return leitor(candidato.curriculo.file).strip()[:LIMITE_CURRICULO]
        finally:
            candidato.curriculo.close()
    except Exception:
        # Arquivo protegido, corrompido, ou PDF que e so imagem escaneada. Nao e
        # erro de programacao e nao deve derrubar a analise.
        logger.warning('Nao consegui ler o curriculo do candidato %s',
                       candidato.pk, exc_info=True)
        return ''


def _perfil_do_candidato(candidato):
    """
    O que o candidato informou, sem os dados que nao ajudam a avaliar.

    Nome, WhatsApp, email e endereco de rua ficam FORA: nao dizem nada sobre
    aptidao e so aumentam a exposicao de dado pessoal a um terceiro. Cidade e
    bairro entram porque deslocamento e criterio real em vaga operacional.
    """
    linhas = []
    if candidato.cidade or candidato.bairro:
        linhas.append(f'Mora em: {candidato.bairro or "?"}, '
                      f'{candidato.cidade or "?"}')
    if candidato.idade_texto:
        # `idade_texto` traz anos E meses. Em vaga de primeiro emprego, a
        # diferenca entre 17 e 11 meses e 18 recem feitos decide contratacao.
        linhas.append(f'Idade: {candidato.idade_texto}')
    if candidato.experiencia_previa:
        linhas.append(f'Experiência prévia: {candidato.experiencia_previa}')
    if candidato.disponibilidade_horario:
        linhas.append(f'Disponibilidade: {candidato.disponibilidade_horario}')

    for chave, valor in (candidato.dados_custom or {}).items():
        linhas.append(f'{chave}: {valor}')

    return '\n'.join(linhas)


def _montar_prompt(candidato, requisitos, texto_curriculo):
    requisitos_texto = '\n'.join(
        f'{i}. {r.texto}' + (' (OBRIGATÓRIO)' if r.obrigatorio else ' (desejável)')
        for i, r in enumerate(requisitos, 1))

    vaga = candidato.vaga.nome_exibido if candidato.vaga_id else 'Banco de talentos'

    sistema = (
        'Você é um analista de RH. Avalia candidatos APENAS contra os '
        'requisitos que a empresa declarou, e nunca inventa critério próprio.\n'
        'Não julgue por gênero, idade, aparência, origem, religião ou classe. '
        'Se um requisito não puder ser verificado com o que foi informado, diga '
        'que falta informação, em vez de supor.\n'
        'Sua saída é SUGESTÃO para um humano decidir. Responda só com JSON.'
    )

    usuario = f"""Vaga: {vaga}

REQUISITOS DECLARADOS PELA EMPRESA:
{requisitos_texto}

O QUE O CANDIDATO INFORMOU:
{_perfil_do_candidato(candidato) or '(não informou nada além do nome e contato)'}

CURRÍCULO:
{texto_curriculo or '(não anexou currículo, ou o arquivo não pôde ser lido)'}

Responda com este JSON, sem texto fora dele:
{{
  "veredito": "apto | atencao | inapto | insuficiente",
  "resumo": "duas ou três linhas sobre o perfil, factuais",
  "sinais_de_atencao": ["ponto a conferir na entrevista", "..."],
  "requisitos": [
    {{"requisito": "texto do requisito", "atende": "sim | nao | nao_da_pra_saber",
      "porque": "uma frase curta baseada no que foi informado"}}
  ]
}}

Use "insuficiente" quando o candidato informou pouco demais para avaliar. Isso
é sobre o DADO, não sobre a pessoa."""

    return [{'role': 'system', 'content': sistema},
            {'role': 'user', 'content': usuario}]


def _extrair_json(resposta):
    """
    Le o JSON da resposta, tolerando cerca de markdown.

    O modelo as vezes devolve ```json ... ``` mesmo mandado nao fazer. Falhar
    por causa disso seria desperdicar uma chamada ja paga.
    """
    if not resposta:
        return None
    texto = resposta.strip()
    if texto.startswith('```'):
        texto = texto.split('```')[1]
        if texto.startswith('json'):
            texto = texto[4:]
    try:
        return json.loads(texto.strip())
    except (ValueError, IndexError):
        logger.warning('Resposta da IA nao era JSON: %.200s', resposta)
        return None


def analisar_candidato(candidato, *, usuario=None):
    """
    Pede a analise e grava o resultado. NAO move o candidato.

    Levanta `TriagemIndisponivel` quando falta integracao de IA ou quando a vaga
    nao tem requisito de triagem cadastrado. O segundo caso e recusa deliberada:
    sem criterio declarado, a IA inventaria o proprio.
    """
    integracao = integracao_ia_do_tenant(candidato.tenant)
    if integracao is None:
        raise TriagemIndisponivel(
            'Nenhuma integração de IA ativa neste tenant. Configure em '
            'Integrações antes de usar a triagem.')

    if not candidato.vaga_id:
        raise TriagemIndisponivel(
            'Candidato do banco de talentos não tem vaga, e portanto não tem '
            'requisitos contra os quais avaliar.')

    requisitos = list(candidato.vaga.requisitos_de_triagem())
    if not requisitos:
        raise TriagemIndisponivel(
            'Esta vaga não tem nenhum requisito marcado para triagem. Cadastre '
            'os critérios na aba Requisitos: sem eles a IA avaliaria por '
            'critério próprio.')

    texto_curriculo = _texto_do_curriculo(candidato)
    mensagens = _montar_prompt(candidato, requisitos, texto_curriculo)

    resposta = chamar_llm(integracao, mensagens, max_tokens=900, timeout=60)
    dados = _extrair_json(resposta)

    if dados is None:
        raise TriagemIndisponivel(
            'A IA não respondeu no formato esperado. Tente de novo; se '
            'persistir, confira a configuração da integração.')

    veredito = dados.get('veredito', '')
    validos = {v for v, _ in AnaliseCandidato.VEREDITO_CHOICES}
    if veredito not in validos:
        # Modelo alucinou uma categoria. Cair em "insuficiente" e o lado seguro:
        # nao afirma nada sobre a pessoa e pede olho humano.
        veredito = AnaliseCandidato.VEREDITO_INSUFICIENTE

    sinais = dados.get('sinais_de_atencao') or []
    if not isinstance(sinais, list):
        sinais = [str(sinais)]

    return AnaliseCandidato.all_tenants.create(
        tenant=candidato.tenant,
        candidato=candidato,
        veredito=veredito,
        resumo=str(dados.get('resumo') or '')[:2000],
        sinais_de_atencao=[str(s)[:300] for s in sinais][:10],
        requisitos_avaliados=dados.get('requisitos') or [],
        modelo=(integracao.configuracoes_extras or {}).get('modelo', '')[:60],
        usou_curriculo=bool(texto_curriculo),
        criado_por=usuario if usuario and usuario.is_authenticated else None,
    )
